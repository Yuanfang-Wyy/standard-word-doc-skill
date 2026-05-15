#!/usr/bin/env python3
"""Audit a DOCX against the standard Word formatting checklist."""

from __future__ import annotations

import argparse
import re
import zipfile
from collections import Counter
from dataclasses import dataclass
from pathlib import Path
from xml.etree import ElementTree as ET

try:
    from docx import Document
except ImportError as exc:  # pragma: no cover
    raise SystemExit("missing dependency: python-docx. Install with: pip install python-docx") from exc


DEFAULT_OUTPUT_DIR = Path.home() / "Documents" / "AI-Stack-Outputs" / "word-docs"
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W_NS}

UNICODE_BULLETS = tuple("•·◆▪●○■□—–")
UNSAFE_FONTS = {"symbol", "wingdings", "wingdings 2", "wingdings 3"}


@dataclass
class Issue:
    rule_id: str
    severity: str
    location: str
    message: str
    action: str
    auto_fixable: bool


def qn(name: str) -> str:
    prefix, tag = name.split(":", 1)
    if prefix != "w":
        raise ValueError(name)
    return f"{{{W_NS}}}{tag}"


def paragraph_text(paragraph) -> str:
    return paragraph.text.strip()


def style_name(paragraph) -> str:
    return (paragraph.style.name if paragraph.style is not None else "").lower()


def is_heading(paragraph) -> bool:
    name = style_name(paragraph)
    return name.startswith("heading") or name.startswith("标题")


def heading_level(paragraph) -> int | None:
    name = style_name(paragraph)
    match = re.search(r"heading\s*([1-4])", name)
    if match:
        return int(match.group(1))
    match = re.search(r"标题\s*([1-4一二三四])", name)
    if match:
        raw = match.group(1)
        return {"一": 1, "二": 2, "三": 3, "四": 4}.get(raw, int(raw) if raw.isdigit() else 0) or None
    return None


def is_body(paragraph) -> bool:
    if is_heading(paragraph):
        return False
    name = style_name(paragraph)
    if name in {"title", "footer"} or "toc" in name or "table" in name:
        return False
    return not name.startswith("list") and paragraph_text(paragraph) != ""


def location_for(index: int, paragraph) -> str:
    text = paragraph_text(paragraph)
    if is_heading(paragraph) and text:
        return f"标题：{text[:60]}"
    if text:
        return f"段落 {index}: {text[:60]}"
    return f"段落 {index}"


def run_size_pt(run) -> float | None:
    if run.font.size is not None:
        return round(float(run.font.size.pt), 1)
    try:
        if run.style and run.style.font.size is not None:
            return round(float(run.style.font.size.pt), 1)
    except Exception:
        return None
    return None


def run_font_names(run) -> set[str]:
    names: set[str] = set()
    if run.font.name:
        names.add(run.font.name.lower())
    rpr = run._element.rPr
    if rpr is not None and rpr.rFonts is not None:
        for value in rpr.rFonts.attrib.values():
            if value:
                names.add(value.lower())
    return names


def has_hard_break(paragraph) -> bool:
    for br in paragraph._p.xpath(".//w:br"):
        if br.get(qn("w:type")) not in {"page", "column"}:
            return True
    return "\v" in paragraph.text


def looks_manual_numbered(text: str) -> bool:
    patterns = [
        r"^\s*\d+(?:\.\d+)*[、.]\s+",
        r"^\s*[（(][一二三四五六七八九十\d]+[）)]",
        r"^\s*第[一二三四五六七八九十百千万\d]+[章节部分篇]",
    ]
    return any(re.search(pattern, text) for pattern in patterns)


def looks_fake_heading(paragraph) -> bool:
    if is_heading(paragraph) or style_name(paragraph) == "title" or not paragraph_text(paragraph):
        return False
    for run in paragraph.runs:
        size = run_size_pt(run)
        if run.bold and size is not None and size >= 14:
            return True
    return False


def has_chinese_english_spacing_issue(text: str) -> bool:
    return bool(re.search(r"[\u4e00-\u9fff][A-Za-z]|[A-Za-z][\u4e00-\u9fff]", text))


def docx_xml(path: Path, name: str) -> str:
    with zipfile.ZipFile(path) as zf:
        if name not in zf.namelist():
            return ""
        return zf.read(name).decode("utf-8", errors="ignore")


def has_percent_table_width(path: Path) -> bool:
    xml = docx_xml(path, "word/document.xml")
    if not xml:
        return False
    try:
        root = ET.fromstring(xml.encode("utf-8"))
    except ET.ParseError:
        return 'w:type="pct"' in xml
    for tbl_w in root.findall(".//w:tblW", NS):
        if tbl_w.get(qn("w:type")) == "pct":
            return True
    return False


def has_page_number_field(path: Path) -> bool:
    with zipfile.ZipFile(path) as zf:
        names = [n for n in zf.namelist() if n == "word/document.xml" or n.startswith("word/footer")]
        for name in names:
            xml = zf.read(name).decode("utf-8", errors="ignore")
            if re.search(r"<w:instrText[^>]*>\s*PAGE\s*</w:instrText>", xml):
                return True
            if re.search(r'w:instr="[^"]*PAGE', xml):
                return True
    return False


def collect_issues(path: Path) -> list[Issue]:
    if not path.exists():
        raise SystemExit(f"input not found: {path}")
    if path.suffix.lower() != ".docx":
        raise SystemExit("only .docx is supported")

    doc = Document(str(path))
    issues: list[Issue] = []
    body_sizes: list[float] = []

    for index, paragraph in enumerate(doc.paragraphs, start=1):
        text = paragraph_text(paragraph)
        loc = location_for(index, paragraph)
        stripped = text.lstrip()

        if stripped.startswith(UNICODE_BULLETS):
            issues.append(Issue("E001", "错误", loc, "段落开头使用 Unicode bullet 字符。", "删除符号并应用 Word List Bullet 样式。", True))

        if text and looks_manual_numbered(text) and not is_heading(paragraph):
            issues.append(Issue("W004", "警告", loc, "正文段落中存在手动编号。", "需人工确认是否改为 Heading 样式或有序列表。", False))

        if looks_fake_heading(paragraph):
            issues.append(Issue("E002", "错误", loc, "使用正文样式加粗和大字号模拟标题。", "按字号映射为 Heading 1-4。", True))

        for run in paragraph.runs:
            unsafe = run_font_names(run) & UNSAFE_FONTS
            if unsafe:
                issues.append(Issue("E003", "错误", loc, f"使用禁止字体：{', '.join(sorted(unsafe))}。", "替换为 Arial/微软雅黑。", True))
            if is_body(paragraph):
                if any("courier" in name for name in run_font_names(run)):
                    continue
                size = run_size_pt(run)
                if size is not None:
                    body_sizes.append(size)

        if has_hard_break(paragraph):
            issues.append(Issue("E004", "错误", loc, "段落中存在硬换行。", "拆分为独立段落。", True))

        if "\u3000" in text:
            issues.append(Issue("S001", "建议", loc, "正文中存在全角空格。", "建议替换为半角空格。", False))

        if text and has_chinese_english_spacing_issue(text):
            issues.append(Issue("S002", "建议", loc, "中英文之间缺少空格。", "建议人工检查，避免误改专有名词。", False))

    unique_sizes = {size for size in body_sizes if size > 0}
    if len(unique_sizes) > 2:
        issues.append(Issue("W001", "警告", "全文", f"正文字号存在多种值：{', '.join(map(str, sorted(unique_sizes)))}pt。", "统一正文为 11pt。", True))

    if has_percent_table_width(path):
        issues.append(Issue("W002", "警告", "表格", "存在百分比表格宽度，WPS/Google Docs 中可能变形。", "转换为 DXA 固定宽度。", True))

    empty_run = 0
    for index, paragraph in enumerate(doc.paragraphs, start=1):
        if paragraph_text(paragraph):
            empty_run = 0
            continue
        empty_run += 1
        if empty_run >= 2:
            issues.append(Issue("W003", "警告", f"段落 {index}", "存在连续空段落。", "保留一个空段落，删除多余空段落。", True))

    previous_level = 0
    for paragraph in doc.paragraphs:
        level = heading_level(paragraph)
        if level is None:
            continue
        if previous_level and level > previous_level + 1:
            issues.append(Issue("W005", "警告", f"标题：{paragraph_text(paragraph)[:60]}", f"标题层级从 H{previous_level} 跳到 H{level}。", "报告给用户确认，不自动修复。", False))
        previous_level = level

    if not has_page_number_field(path):
        issues.append(Issue("S003", "建议", "页脚", "文档未检测到页码字段。", "建议添加页码字段或刷新模板页脚。", False))

    return issues


def summarize(issues: list[Issue]) -> Counter:
    return Counter(issue.severity for issue in issues)


def write_report(path: Path, issues: list[Issue], output: Path) -> None:
    output.parent.mkdir(parents=True, exist_ok=True)
    counts = summarize(issues)
    with output.open("w", encoding="utf-8") as fh:
        fh.write("# Word 格式审计报告\n\n")
        fh.write(f"- 文件：`{path}`\n")
        fh.write(f"- 问题总数：{len(issues)}\n")
        fh.write(f"- 错误：{counts.get('错误', 0)}\n")
        fh.write(f"- 警告：{counts.get('警告', 0)}\n")
        fh.write(f"- 建议：{counts.get('建议', 0)}\n\n")

        for severity in ("错误", "警告", "建议"):
            fh.write(f"## {severity}\n\n")
            filtered = [issue for issue in issues if issue.severity == severity]
            if not filtered:
                fh.write("- 无。\n\n")
                continue
            for issue in filtered:
                fix = "可自动修复" if issue.auto_fixable else "需人工确认"
                fh.write(f"- `{issue.rule_id}` {issue.location}：{issue.message} {issue.action}（{fix}）\n")
            fh.write("\n")


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("input", type=Path)
    parser.add_argument("--output", type=Path)
    args = parser.parse_args()

    input_path = args.input.expanduser()
    output = args.output.expanduser() if args.output else DEFAULT_OUTPUT_DIR / f"{input_path.stem}_audit.md"
    issues = collect_issues(input_path)
    write_report(input_path, issues, output)
    print(output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
