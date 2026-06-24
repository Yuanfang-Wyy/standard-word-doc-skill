#!/usr/bin/env python3
"""Render a formal DOCX from Markdown using python-docx.

The script prefers assets/standard-word-template.docx when it exists. If the
template is missing, it creates a new document with the default style
specification from references/style-guide.md.
"""

from __future__ import annotations

import argparse
import re
import uuid
from dataclasses import dataclass
from datetime import date
from pathlib import Path

try:
    from docx import Document
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor
except ImportError as exc:  # pragma: no cover - exercised by dependency checks
    raise SystemExit("missing dependency: python-docx. Install with: pip install python-docx") from exc


SKILL_DIR = Path(__file__).resolve().parents[1]
TEMPLATE = SKILL_DIR / "assets" / "standard-word-template.docx"
DEFAULT_OUTPUT_DIR = Path.home() / "Documents" / "AI-Stack-Outputs" / "word-docs"

WESTERN_FONT = "Times New Roman"
EAST_ASIA_FONT = "仿宋"
FORBIDDEN_BULLET_PREFIX = "•·◆▪●○■□—–"

STYLE_SPEC = {
    "Title": {"size": 28, "bold": True, "color": "000000", "before": 0, "after": 4, "line": 1.0},
    "Heading 1": {"size": 22, "bold": True, "color": "000000", "before": 24, "after": 4, "line": 1.5},
    "Heading 2": {"size": 16, "bold": True, "color": "000000", "before": 8, "after": 4, "line": 1.5},
    "Heading 3": {"size": 15, "bold": True, "color": "000000", "before": 8, "after": 4, "line": 1.5},
    "Heading 4": {"size": 14, "bold": True, "color": "000000", "before": 4, "after": 2, "line": 1.5},
    "Normal": {"size": 11, "bold": False, "color": "000000", "before": 0, "after": 8, "line": 1.15},
    "Body Text": {"size": 11, "bold": False, "color": "000000", "before": 0, "after": 8, "line": 1.15},
    "List Bullet": {"size": 11, "bold": False, "color": "000000", "before": 0, "after": 6, "line": 1.15},
    "List Number": {"size": 11, "bold": False, "color": "000000", "before": 0, "after": 6, "line": 1.15},
}

STANDARD_TABLE = {
    "header_fill": "1F5FAE",
    "odd_fill": "F3F6FB",
    "even_fill": "FFFFFF",
    "border": "C9C9C9",
    "header_color": "FFFFFF",
    "body_color": "000000",
    "font_size": 12,
}


@dataclass
class Block:
    kind: str
    text: str = ""
    level: int = 0
    rows: list[list[str]] | None = None
    lang: str = ""
    indent: int = 0


def parse_frontmatter(text: str) -> tuple[dict[str, str], str]:
    stripped = text.lstrip("\ufeff")
    if not stripped.startswith("---\n"):
        return {}, text
    end = stripped.find("\n---", 4)
    if end == -1:
        return {}, text
    raw = stripped[4:end].strip()
    body = stripped[stripped.find("\n", end + 1) + 1 :]
    meta: dict[str, str] = {}
    for line in raw.splitlines():
        if ":" not in line:
            continue
        key, value = line.split(":", 1)
        meta[key.strip()] = value.strip().strip('"').strip("'")
    return meta, body


def split_table(lines: list[str], start: int) -> tuple[Block | None, int]:
    if start + 1 >= len(lines):
        return None, start
    if "|" not in lines[start] or "|" not in lines[start + 1]:
        return None, start
    separator = [c.strip() for c in lines[start + 1].strip().strip("|").split("|")]
    if not separator or not all(re.fullmatch(r":?-{3,}:?", c or "") for c in separator):
        return None, start
    rows: list[list[str]] = []
    i = start
    while i < len(lines) and "|" in lines[i]:
        rows.append([cleanup_inline(c.strip()) for c in lines[i].strip().strip("|").split("|")])
        i += 1
    if len(rows) >= 2:
        rows.pop(1)
    max_cols = max((len(row) for row in rows), default=0)
    if max_cols:
        rows = [row + [""] * (max_cols - len(row)) for row in rows]
    return Block(kind="table", rows=rows), i


def cleanup_inline(text: str) -> str:
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    return text.strip()


def strip_heading_number(text: str) -> str:
    patterns = [
        r"^第[一二三四五六七八九十百千万\d]+[章节部分篇][、.\s-]*",
        r"^\d+(?:\.\d+)*[、.\s-]+",
        r"^[（(][一二三四五六七八九十\d]+[）)]\s*",
    ]
    cleaned = text.strip()
    for pattern in patterns:
        cleaned = re.sub(pattern, "", cleaned)
    return cleaned.strip() or text.strip()


def parse_markdown(text: str) -> tuple[dict[str, str], list[Block]]:
    meta, body = parse_frontmatter(text)
    lines = body.splitlines()
    blocks: list[Block] = []
    paragraph: list[str] = []
    in_code = False
    code_lang = ""
    code_lines: list[str] = []

    def flush_paragraph() -> None:
        nonlocal paragraph
        if paragraph:
            blocks.append(Block(kind="paragraph", text=cleanup_inline(" ".join(paragraph).strip())))
            paragraph = []

    i = 0
    while i < len(lines):
        raw = lines[i].rstrip("\n")
        line = raw.strip()

        fence = re.match(r"^```([A-Za-z0-9_-]*)\s*$", line)
        if fence:
            if in_code:
                blocks.append(Block(kind="code", text="\n".join(code_lines), lang=code_lang))
                code_lines = []
                code_lang = ""
                in_code = False
            else:
                flush_paragraph()
                code_lang = fence.group(1)
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(raw)
            i += 1
            continue

        if not line:
            flush_paragraph()
            i += 1
            continue

        table, new_i = split_table(lines, i)
        if table:
            flush_paragraph()
            blocks.append(table)
            i = new_i
            continue

        heading = re.match(r"^(#{1,4})\s+(.+)$", line)
        if heading:
            flush_paragraph()
            blocks.append(Block(kind="heading", level=len(heading.group(1)), text=strip_heading_number(cleanup_inline(heading.group(2)))))
            i += 1
            continue

        bullet = re.match(r"^(\s*)([-*+])\s+(.+)$", raw)
        unicode_bullet = re.match(rf"^(\s*)[{re.escape(FORBIDDEN_BULLET_PREFIX)}]\s*(.+)$", raw)
        if bullet or unicode_bullet:
            flush_paragraph()
            indent = len((bullet or unicode_bullet).group(1).replace("\t", "    ")) // 2
            text_value = (bullet.group(3) if bullet else unicode_bullet.group(2)).strip()  # type: ignore[union-attr]
            blocks.append(Block(kind="bullet", text=cleanup_inline(text_value), indent=indent))
            i += 1
            continue

        ordered = re.match(r"^(\s*)\d+[.)]\s+(.+)$", raw)
        if ordered:
            flush_paragraph()
            indent = len(ordered.group(1).replace("\t", "    ")) // 2
            blocks.append(Block(kind="number", text=cleanup_inline(ordered.group(2).strip()), indent=indent))
            i += 1
            continue

        paragraph.append(line)
        i += 1

    if in_code:
        blocks.append(Block(kind="code", text="\n".join(code_lines), lang=code_lang))
    flush_paragraph()
    return meta, blocks


def document_title(meta: dict[str, str], blocks: list[Block]) -> str:
    for key in ("title", "project_name", "项目名称"):
        if meta.get(key):
            return meta[key]
    for block in blocks:
        if block.kind == "heading" and block.level == 1:
            return block.text
    return "标准文档"


def document_type(meta: dict[str, str], blocks: list[Block]) -> str:
    for key in ("document_type", "doc_type", "文档类型"):
        if meta.get(key):
            return meta[key]
    return "正式文档"


def sanitize_filename(text: str) -> str:
    text = re.sub(r"[\\/:*?\"<>|]+", "_", text.strip())
    text = re.sub(r"\s+", "_", text)
    return text[:80] or f"standard-doc-{uuid.uuid4().hex[:8]}"


def hex_to_rgb(value: str) -> RGBColor:
    value = value.strip().lstrip("#")
    return RGBColor(int(value[0:2], 16), int(value[2:4], 16), int(value[4:6], 16))


def set_font_east_asia(obj, western: str = WESTERN_FONT, east_asia: str = EAST_ASIA_FONT) -> None:
    rpr = obj._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr, value in (
        ("w:ascii", western),
        ("w:hAnsi", western),
        ("w:cs", western),
        ("w:eastAsia", east_asia),
    ):
        rfonts.set(qn(attr), value)


def apply_run_format(run, size: float, bold: bool, color: str, *, monospace: bool = False) -> None:
    run.font.name = "Courier New" if monospace else WESTERN_FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = hex_to_rgb(color)
    set_font_east_asia(run, western="Courier New" if monospace else WESTERN_FONT)


def apply_paragraph_format(paragraph, before: float, after: float, line: float) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def style_if_exists(doc, style_name: str):
    try:
        return doc.styles[style_name]
    except KeyError:
        return None


def set_style_font(style, size: float, bold: bool, color: str) -> None:
    style.font.name = WESTERN_FONT
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = hex_to_rgb(color)
    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr, value in (
        ("w:ascii", WESTERN_FONT),
        ("w:hAnsi", WESTERN_FONT),
        ("w:cs", WESTERN_FONT),
        ("w:eastAsia", EAST_ASIA_FONT),
    ):
        rfonts.set(qn(attr), value)


def ensure_default_styles(doc) -> None:
    for style_name, spec in STYLE_SPEC.items():
        style = style_if_exists(doc, style_name)
        if style is None:
            continue
        set_style_font(style, spec["size"], spec["bold"], spec["color"])
        if hasattr(style, "paragraph_format"):
            style.paragraph_format.space_before = Pt(spec["before"])
            style.paragraph_format.space_after = Pt(spec["after"])
            style.paragraph_format.line_spacing = spec["line"]


def configure_page(doc) -> None:
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)


def clear_document_body(doc) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_paragraph_style(paragraph, style_name: str) -> None:
    try:
        paragraph.style = style_name
    except KeyError:
        paragraph.style = "Normal"


def add_text_paragraph(doc, text: str, style_name: str = "Normal", *, align=None, direct_format: bool = True):
    paragraph = doc.add_paragraph()
    set_paragraph_style(paragraph, style_name)
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    if direct_format:
        spec = STYLE_SPEC.get(style_name, STYLE_SPEC["Normal"])
        apply_paragraph_format(paragraph, spec["before"], spec["after"], spec["line"])
        apply_run_format(run, spec["size"], spec["bold"], spec["color"])
    return paragraph


def add_heading(doc, text: str, level: int, *, direct_format: bool = True) -> None:
    style_name = f"Heading {min(max(level, 1), 4)}"
    add_text_paragraph(doc, strip_heading_number(text), style_name, direct_format=direct_format)


def add_list_item(doc, text: str, style_name: str, indent: int, *, direct_format: bool = True) -> None:
    paragraph = add_text_paragraph(doc, text, style_name, direct_format=direct_format)
    if indent:
        paragraph.paragraph_format.left_indent = Cm(0.75 * indent)


def add_field_run(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, text, fld_end])


def add_page_number_footer(doc) -> None:
    for section in doc.sections:
        paragraph = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_style(paragraph, "Footer")
        paragraph.clear()
        run = paragraph.add_run("第 ")
        apply_run_format(run, 9, False, "808080")
        add_field_run(paragraph, "PAGE")
        run = paragraph.add_run(" 页")
        apply_run_format(run, 9, False, "808080")


def add_toc_placeholder(doc, *, direct_format: bool = True) -> None:
    title = add_text_paragraph(doc, "目 录", "Heading 1", align=WD_ALIGN_PARAGRAPH.CENTER, direct_format=direct_format)
    title.paragraph_format.space_before = Pt(0)
    paragraph = doc.add_paragraph()
    set_paragraph_style(paragraph, "Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_field_run(paragraph, r'TOC \o "1-4" \h \z \u')
    note = doc.add_paragraph("目录字段将在 Word/WPS 中刷新后显示准确页码。")
    set_paragraph_style(note, "Normal")
    doc.add_page_break()


def add_cover(doc, meta: dict[str, str], blocks: list[Block], *, direct_format: bool = True) -> None:
    title = document_title(meta, blocks)
    doc_type = document_type(meta, blocks)
    organization = meta.get("organization") or meta.get("编制单位") or "XXXX单位"
    classification = meta.get("classification") or meta.get("文件密级") or "内部资料"
    version = meta.get("version") or meta.get("版本") or "V1.0"
    doc_date = meta.get("date") or meta.get("日期") or f"{date.today():%Y年%m月%d日}"

    add_text_paragraph(doc, title, "Title", align=WD_ALIGN_PARAGRAPH.CENTER, direct_format=direct_format)
    add_text_paragraph(doc, doc_type, "Title", align=WD_ALIGN_PARAGRAPH.CENTER, direct_format=direct_format)
    doc.add_paragraph()
    for line in (
        f"文件密级：{classification}",
        f"编制单位：{organization}",
        f"编制日期：{doc_date}",
        f"版本号：{version}",
    ):
        add_text_paragraph(doc, line, "Normal", align=WD_ALIGN_PARAGRAPH.CENTER, direct_format=direct_format)
    doc.add_page_break()


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = STANDARD_TABLE["border"], size: str = "8") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, width_emu: int) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(int(width_emu / 635)))


def apply_cell_text(cell, text: str, *, header: bool) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if header else WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    apply_paragraph_format(paragraph, 4, 4, 1.15)
    apply_run_format(
        run,
        STANDARD_TABLE["font_size"],
        True,
        STANDARD_TABLE["header_color"] if header else STANDARD_TABLE["body_color"],
    )


def add_table(doc, rows: list[list[str]], *, direct_format: bool = True) -> None:
    if not rows:
        return
    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    section = doc.sections[0]
    content_width = int(section.page_width - section.left_margin - section.right_margin)
    set_table_width(table, content_width)
    col_width = int(content_width / cols)

    for row_index, row in enumerate(rows):
        for col_index in range(cols):
            cell = table.cell(row_index, col_index)
            cell.width = col_width
            text = row[col_index] if col_index < len(row) else ""
            is_header = row_index == 0
            if is_header:
                fill = STANDARD_TABLE["header_fill"]
            else:
                fill = STANDARD_TABLE["odd_fill"] if row_index % 2 == 1 else STANDARD_TABLE["even_fill"]
            set_cell_shading(cell, fill)
            set_cell_border(cell)
            apply_cell_text(cell, text, header=is_header)
    doc.add_paragraph()


def add_code_block(doc, text: str, lang: str = "", *, direct_format: bool = True) -> None:
    if lang:
        add_text_paragraph(doc, f"代码示例（{lang}）", "Normal", direct_format=direct_format)
    for line in text.splitlines() or [""]:
        paragraph = doc.add_paragraph()
        set_paragraph_style(paragraph, "Normal")
        paragraph.paragraph_format.left_indent = Cm(0.3)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        run = paragraph.add_run(line)
        if direct_format:
            apply_run_format(run, 9, False, "404040", monospace=True)


def render(input_path: Path, output_dir: Path, filename: str | None, template_path: Path = TEMPLATE) -> Path:
    if not input_path.exists():
        raise SystemExit(f"input not found: {input_path}")
    meta, blocks = parse_markdown(input_path.read_text(encoding="utf-8"))
    if not blocks:
        raise SystemExit("input markdown has no content")

    use_template = template_path.exists()
    doc = Document(str(template_path)) if use_template else Document()
    clear_document_body(doc)
    if not use_template:
        configure_page(doc)
        ensure_default_styles(doc)
        add_page_number_footer(doc)

    add_cover(doc, meta, blocks, direct_format=not use_template)
    add_toc_placeholder(doc, direct_format=not use_template)

    title = document_title(meta, blocks)
    skipped_title = False
    for block in blocks:
        if block.kind == "heading":
            if block.level == 1 and not skipped_title and block.text == title:
                skipped_title = True
                continue
            add_heading(doc, block.text, block.level, direct_format=not use_template)
        elif block.kind == "paragraph":
            add_text_paragraph(doc, block.text, "Body Ref" if use_template and style_if_exists(doc, "Body Ref") else "Normal", direct_format=not use_template)
        elif block.kind == "bullet":
            list_style = "List Paragraph" if use_template and style_if_exists(doc, "List Paragraph") else "List Bullet"
            add_list_item(doc, block.text, list_style, block.indent, direct_format=not use_template)
        elif block.kind == "number":
            list_style = "List Paragraph" if use_template and style_if_exists(doc, "List Paragraph") else "List Number"
            add_list_item(doc, block.text, list_style, block.indent, direct_format=not use_template)
        elif block.kind == "table" and block.rows:
            add_table(doc, block.rows, direct_format=not use_template)
        elif block.kind == "code":
            add_code_block(doc, block.text, block.lang, direct_format=not use_template)

    output_dir.mkdir(parents=True, exist_ok=True)
    out_name = filename or f"{sanitize_filename(document_title(meta, blocks))}.docx"
    if not out_name.lower().endswith(".docx"):
        out_name += ".docx"
    output_path = output_dir / out_name
    doc.save(output_path)
    return output_path


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("input", type=Path, help="Markdown input")
    parser.add_argument("--output-dir", type=Path, default=DEFAULT_OUTPUT_DIR)
    parser.add_argument("--filename", help="Output file name")
    parser.add_argument("--template", type=Path, default=TEMPLATE, help="Optional DOCX template to inherit")
    args = parser.parse_args()
    out = render(args.input, args.output_dir.expanduser(), args.filename, args.template.expanduser())
    print(out)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
