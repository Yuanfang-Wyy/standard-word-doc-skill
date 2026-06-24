#!/usr/bin/env python3
"""Repair a DOCX according to the standard Word formatting checklist."""

from __future__ import annotations

import argparse
import re
from collections import Counter
from pathlib import Path

try:
    from docx import Document
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.table import Table
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor
    from docx.text.paragraph import Paragraph
except ImportError as exc:  # pragma: no cover
    raise SystemExit("missing dependency: python-docx. Install with: pip install python-docx") from exc

from audit_standard_docx import (
    DEFAULT_OUTPUT_DIR,
    UNICODE_BULLETS,
    UNSAFE_FONTS,
    collect_issues,
    heading_level,
    is_body,
    is_heading,
    looks_fake_heading,
    paragraph_text,
    run_font_names,
    run_size_pt,
    style_name,
)


SKILL_DIR = Path(__file__).resolve().parents[1]
TEMPLATE = SKILL_DIR / "assets" / "standard-word-template.docx"
WESTERN_FONT = "Times New Roman"
EAST_ASIA_FONT = "仿宋"
PARALLEL_BULLET = "◆"
STANDARD_TABLE = {
    "header_fill": "1F5FAE",
    "odd_fill": "F3F6FB",
    "even_fill": "FFFFFF",
    "border": "C9C9C9",
    "header_color": "FFFFFF",
    "body_color": "000000",
    "font_size": 12,
}


def set_font_east_asia(run, western: str = WESTERN_FONT, east_asia: str = EAST_ASIA_FONT) -> None:
    run.font.name = western
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr, value in (
        ("w:ascii", western),
        ("w:hAnsi", western),
        ("w:cs", western),
        ("w:eastAsia", east_asia),
    ):
        rfonts.set(qn(attr), value)


def hex_to_rgb(value: str) -> RGBColor:
    value = value.strip().lstrip("#")
    return RGBColor(int(value[0:2], 16), int(value[2:4], 16), int(value[4:6], 16))


def safe_set_style(paragraph, style_name: str) -> None:
    try:
        paragraph.style = style_name
    except KeyError:
        paragraph.style = "Normal"


def style_exists(doc, style_name: str) -> bool:
    try:
        doc.styles[style_name]
        return True
    except KeyError:
        return False


def body_style_name(doc) -> str:
    return "Body Ref" if style_exists(doc, "Body Ref") else "Normal"


def list_style_name(doc) -> str:
    if style_exists(doc, "List Paragraph"):
        return "List Paragraph"
    if style_exists(doc, "List Bullet"):
        return "List Bullet"
    return body_style_name(doc)


def clear_document_body(doc) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def rewrite_paragraph(paragraph, text: str, style_name: str | None = None) -> None:
    paragraph.clear()
    if style_name:
        safe_set_style(paragraph, style_name)
    run = paragraph.add_run(text)
    set_font_east_asia(run)
    run.font.size = Pt(11)


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def insert_paragraph_after(paragraph, text: str) -> Paragraph:
    new_element = OxmlElement("w:p")
    paragraph._p.addnext(new_element)
    new_paragraph = Paragraph(new_element, paragraph._parent)
    try:
        new_paragraph.style = paragraph.style
    except Exception:
        safe_set_style(new_paragraph, "Normal")
    run = new_paragraph.add_run(text)
    set_font_east_asia(run)
    run.font.size = Pt(11)
    return new_paragraph


def paragraph_has_hard_break(paragraph) -> bool:
    for br in paragraph._p.xpath(".//w:br"):
        if br.get(qn("w:type")) not in {"page", "column"}:
            return True
    return "\v" in paragraph.text or "\n" in paragraph.text


def repair_unicode_bullet(paragraph) -> bool:
    text = paragraph_text(paragraph)
    stripped = text.lstrip()
    if not stripped.startswith(UNICODE_BULLETS):
        return False
    leading_spaces = len(text) - len(stripped)
    cleaned = stripped.lstrip("".join(UNICODE_BULLETS)).strip()
    rewrite_paragraph(paragraph, " " * leading_spaces + cleaned, "List Bullet")
    return True


def repair_fake_heading(paragraph) -> bool:
    if not looks_fake_heading(paragraph):
        return False
    max_size = max((run_size_pt(run) or 0 for run in paragraph.runs), default=0)
    if max_size >= 18:
        style = "Heading 1"
    elif max_size >= 14:
        style = "Heading 2"
    elif max_size >= 12:
        style = "Heading 3"
    else:
        style = "Heading 4"
    safe_set_style(paragraph, style)
    for run in paragraph.runs:
        set_font_east_asia(run)
    return True


def repair_unsafe_fonts(paragraph) -> int:
    count = 0
    for run in paragraph.runs:
        if run_font_names(run) & UNSAFE_FONTS:
            set_font_east_asia(run)
            count += 1
    return count


def repair_hard_break(paragraph) -> bool:
    if not paragraph_has_hard_break(paragraph):
        return False
    parts = [part.strip() for part in re.split(r"[\n\v]+", paragraph.text) if part.strip()]
    if len(parts) <= 1:
        return False
    rewrite_paragraph(paragraph, parts[0])
    cursor = paragraph
    for part in parts[1:]:
        cursor = insert_paragraph_after(cursor, part)
    return True


def repair_body_font_sizes(doc) -> int:
    body_sizes: set[float] = set()
    for paragraph in doc.paragraphs:
        if not is_body(paragraph):
            continue
        for run in paragraph.runs:
            size = run_size_pt(run)
            if size is not None:
                body_sizes.add(size)
    if len(body_sizes) <= 2:
        return 0
    changed = 0
    for paragraph in doc.paragraphs:
        if not is_body(paragraph):
            continue
        for run in paragraph.runs:
            run.font.size = Pt(11)
            set_font_east_asia(run)
            changed += 1
    return changed


def set_table_width(table, width_emu: int) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(int(width_emu / 635)))


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = STANDARD_TABLE["border"], size: str = "8") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def apply_standard_cell_text(cell, text: str, *, header: bool) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if header else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.left_indent = None
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(text)
    set_font_east_asia(run)
    run.font.size = Pt(STANDARD_TABLE["font_size"])
    run.font.bold = header
    run.font.color.rgb = hex_to_rgb(STANDARD_TABLE["header_color"] if header else STANDARD_TABLE["body_color"])


def repair_tables(doc) -> int:
    if not doc.tables:
        return 0
    section = doc.sections[0]
    content_width = int(section.page_width - section.left_margin - section.right_margin)
    changed = 0
    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        set_table_width(table, content_width)
        cols = len(table.columns) or 1
        col_width = int(content_width / cols)
        for row in table.rows:
            for cell in row.cells:
                cell.width = col_width
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_font_east_asia(run)
        changed += 1
    return changed


def repair_empty_paragraphs(doc) -> int:
    removed = 0
    empty_run = 0
    for paragraph in list(doc.paragraphs):
        if paragraph_text(paragraph):
            empty_run = 0
            continue
        empty_run += 1
        if empty_run >= 2:
            remove_paragraph(paragraph)
            removed += 1
    return removed


def iter_body_blocks(doc):
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, doc._body)
        elif child.tag == qn("w:tbl"):
            yield Table(child, doc._body)


def max_run_size(paragraph) -> float:
    return max((run_size_pt(run) or 0 for run in paragraph.runs), default=0)


def strip_heading_number(text: str) -> str:
    patterns = [
        r"^第[一二三四五六七八九十百千万\d]+[章节部分篇][、.\s-]*",
        r"^[一二三四五六七八九十]+[、.．]\s*",
        r"^\d+(?:\.\d+)*[、.．\s-]+",
        r"^[（(][一二三四五六七八九十\d]+[）)]\s*",
    ]
    cleaned = text.strip()
    for pattern in patterns:
        cleaned = re.sub(pattern, "", cleaned)
    return cleaned.strip() or text.strip()


def mapped_style_name(source_paragraph, target_doc) -> str:
    text = paragraph_text(source_paragraph)
    if not text:
        return body_style_name(target_doc)

    source_level = heading_level(source_paragraph)
    if source_level is not None:
        heading_style = f"Heading {source_level}"
        return heading_style if style_exists(target_doc, heading_style) else body_style_name(target_doc)

    if looks_fake_heading(source_paragraph):
        size = max_run_size(source_paragraph)
        if size >= 22 and style_exists(target_doc, "Title"):
            return "Title"
        if size >= 18 and style_exists(target_doc, "Heading 1"):
            return "Heading 1"
        if size >= 14 and style_exists(target_doc, "Heading 2"):
            return "Heading 2"
        if size >= 12 and style_exists(target_doc, "Heading 3"):
            return "Heading 3"

    style = style_name(source_paragraph)
    stripped = text.lstrip()
    if "list" in style or "列表" in style or stripped.startswith(UNICODE_BULLETS) or stripped.startswith(PARALLEL_BULLET):
        return list_style_name(target_doc)

    if max_run_size(source_paragraph) >= 24 and len(text) <= 80 and style_exists(target_doc, "Title"):
        return "Title"

    return body_style_name(target_doc)


def add_template_paragraph(target_doc, text: str, style_name: str) -> None:
    paragraph = target_doc.add_paragraph()
    safe_set_style(paragraph, style_name)
    bullet_chars = "".join(UNICODE_BULLETS) + PARALLEL_BULLET
    stripped = text.lstrip()
    cleaned = stripped.lstrip(bullet_chars).strip() if stripped.startswith(tuple(bullet_chars)) else text
    if style_name.lower().startswith("heading"):
        cleaned = strip_heading_number(cleaned)
    if "list" in style_name.lower() or "列表" in style_name:
        cleaned = f"{PARALLEL_BULLET}{cleaned.lstrip(PARALLEL_BULLET).strip()}"
    paragraph.add_run(cleaned)


def add_template_table(target_doc, source_table) -> None:
    rows = source_table.rows
    if not rows:
        return
    col_count = max((len(row.cells) for row in rows), default=0)
    if col_count == 0:
        return

    table = target_doc.add_table(rows=len(rows), cols=col_count)
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    section = target_doc.sections[0]
    content_width = int(section.page_width - section.left_margin - section.right_margin)
    set_table_width(table, content_width)
    col_width = int(content_width / col_count)

    for row_index, row in enumerate(rows):
        for col_index in range(col_count):
            target_cell = table.cell(row_index, col_index)
            target_cell.width = col_width
            source_text = ""
            if col_index < len(row.cells):
                source_text = "\n".join(p.text.strip() for p in row.cells[col_index].paragraphs if p.text.strip())
            is_header = row_index == 0
            if is_header:
                fill = STANDARD_TABLE["header_fill"]
            else:
                fill = STANDARD_TABLE["odd_fill"] if row_index % 2 == 1 else STANDARD_TABLE["even_fill"]
            set_cell_shading(target_cell, fill)
            set_cell_border(target_cell)
            apply_standard_cell_text(target_cell, source_text, header=is_header)


def rebuild_from_template(input_path: Path, output_path: Path) -> None:
    source_doc = Document(str(input_path))
    target_doc = Document(str(TEMPLATE))
    clear_document_body(target_doc)

    for block in iter_body_blocks(source_doc):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if not text:
                continue
            add_template_paragraph(target_doc, text, mapped_style_name(block, target_doc))
        elif isinstance(block, Table):
            add_template_table(target_doc, block)

    target_doc.save(output_path)


def output_paths(input_path: Path, output: Path | None) -> tuple[Path, Path]:
    if output:
        repaired = output
    else:
        repaired = input_path.with_name(f"{input_path.stem}_repaired.docx")
    if repaired.resolve() == input_path.resolve():
        raise SystemExit("repair output must not overwrite the original file")
    summary = repaired.with_name(f"{repaired.stem.replace('_repaired', '')}_repair_summary.md")
    return repaired, summary


def write_summary(input_path: Path, repaired_path: Path, summary_path: Path, before, after, fixes: Counter) -> None:
    unresolved = [issue for issue in after if not issue.auto_fixable]
    with summary_path.open("w", encoding="utf-8") as fh:
        fh.write("# Word 格式修复摘要\n\n")
        fh.write(f"- 原文件：`{input_path}`\n")
        fh.write(f"- 修复后文件：`{repaired_path}`\n")
        fh.write(f"- 修复前问题数：{len(before)}\n")
        fh.write(f"- 修复后剩余问题数：{len(after)}\n\n")
        fh.write("## 自动修复\n\n")
        if fixes:
            for rule_id, count in sorted(fixes.items()):
                fh.write(f"- `{rule_id}`：{count} 处\n")
        else:
            fh.write("- 未执行自动修复。\n")
        fh.write("\n## 仍需人工确认\n\n")
        if unresolved:
            for issue in unresolved:
                fh.write(f"- `{issue.rule_id}` {issue.location}：{issue.message} {issue.action}\n")
        else:
            fh.write("- 无。\n")


def repair(input_path: Path, output: Path | None = None) -> tuple[Path, Path]:
    if not input_path.exists():
        raise SystemExit(f"input not found: {input_path}")
    if input_path.suffix.lower() != ".docx":
        raise SystemExit("only .docx is supported")

    before = collect_issues(input_path)
    repaired_path, summary_path = output_paths(input_path, output)
    repaired_path.parent.mkdir(parents=True, exist_ok=True)

    if TEMPLATE.exists():
        rebuild_from_template(input_path, repaired_path)
        after = collect_issues(repaired_path)
        write_summary(input_path, repaired_path, summary_path, before, after, Counter({"TEMPLATE_REBUILD": 1}))
        return repaired_path, summary_path

    before_rule_ids = {issue.rule_id for issue in before}
    doc = Document(str(input_path))
    fixes: Counter = Counter()

    for paragraph in list(doc.paragraphs):
        if repair_unicode_bullet(paragraph):
            fixes["E001"] += 1
        if repair_fake_heading(paragraph):
            fixes["E002"] += 1
        unsafe_count = repair_unsafe_fonts(paragraph)
        if unsafe_count:
            fixes["E003"] += unsafe_count
        if repair_hard_break(paragraph):
            fixes["E004"] += 1

    size_changes = repair_body_font_sizes(doc)
    if size_changes:
        fixes["W001"] += size_changes

    table_changes = repair_tables(doc)
    if table_changes and "W002" in before_rule_ids:
        fixes["W002"] += table_changes

    empty_changes = repair_empty_paragraphs(doc)
    if empty_changes:
        fixes["W003"] += empty_changes

    doc.save(repaired_path)
    after = collect_issues(repaired_path)
    write_summary(input_path, repaired_path, summary_path, before, after, fixes)
    return repaired_path, summary_path


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("input", type=Path)
    parser.add_argument("--output", type=Path, help="Optional repaired .docx output path")
    args = parser.parse_args()
    repaired, summary = repair(args.input.expanduser(), args.output.expanduser() if args.output else None)
    print(repaired)
    print(summary)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
